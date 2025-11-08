import os
import sys
import uuid
import time

# Asegurar que el paquete app esté en sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from app.training_data_loader import training_loader
from app.rag import rag_engine
from docx import Document


def test_docx_indexing_creates_metadata_source():
    """Crea un .docx temporal en app/documents, ejecuta la carga y verifica que el documento fue indexado en ChromaDB con su 'source'."""
    docs_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'app', 'documents'))
    os.makedirs(docs_dir, exist_ok=True)

    fname = f"test_sample_{uuid.uuid4().hex}.docx"
    path = os.path.join(docs_dir, fname)

    # Crear documento .docx simple
    doc = Document()
    doc.add_heading('Sección de Prueba', level=1)
    doc.add_paragraph('Este es un párrafo de prueba para indexación.')
    doc.save(path)

    try:
        # Ejecutar la carga (síncrona)
        ok = training_loader.load_all_training_data()
        assert ok is True

        # Consultar la colección y buscar metadatas con source == fname
        res = rag_engine.collection.get(include=['metadatas','documents'], limit=1000)
        metas = res.get('metadatas', []) or []

        found = any(m.get('source') == fname for m in metas)
        # Si no se encuentra exactamente por nombre, también intentar buscar por filename parcial en documentos
        if not found:
            docs = res.get('documents', []) or []
            found = any(fname.replace('.docx','') in (d or '') for d in docs)

        assert found, f"El documento {fname} no fue indexado. Muestras de metadatas: {metas[:5]}"

    finally:
        # Limpiar el archivo creado
        try:
            os.remove(path)
        except Exception:
            pass
