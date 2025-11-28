# training_data_loader.py - VERSIÓN FINAL, CORREGIDA Y OPTIMIZADA
import json
import os
import glob
import logging
import re
from typing import List, Dict, Any
from datetime import datetime
from app.rag import rag_engine  # ← IMPORTA rag_engine (inicializado después de chroma_config)

# Soporte para documentos Word
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx no instalado. No se procesarán .docx")

# NUEVO: Importar chunker inteligente
try:
    from app.intelligent_chunker import semantic_chunker
    INTELLIGENT_CHUNKER_AVAILABLE = True
    logging.info("✅ Chunker inteligente disponible")
except ImportError:
    INTELLIGENT_CHUNKER_AVAILABLE = False
    logging.warning("⚠️ Chunker inteligente no disponible")

# Soporte para documentos PDF
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("pdfplumber no instalado. No se procesarán .pdf")

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Procesa documentos Word de Duoc UC para RAG"""
    
    def __init__(self):
        self.processed_count = 0
        logger.info("DocumentProcessor inicializado")

    def extract_from_docx(self, file_path: str) -> List[Dict[str, str]]:
        """Extrae contenido de DOCX usando chunking inteligente si está disponible"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx no disponible")
            return []

        filename = os.path.basename(file_path)
        
        # NUEVO: Usar chunker inteligente si está disponible
        if INTELLIGENT_CHUNKER_AVAILABLE:
            logger.info(f"📄 Procesando con CHUNKER INTELIGENTE: {filename}")
            try:
                category = self._detect_category_from_filename(filename)
                chunks = semantic_chunker.chunk_document_from_path(file_path, filename, category)
                
                # Convertir chunks a formato esperado
                result = []
                for chunk in chunks:
                    result.append({
                        'text': chunk.content,
                        'section': chunk.section,
                        'style': 'SemanticChunk',
                        'is_structured': True,
                        'page_reference': chunk.chunk_id,
                        'keywords': chunk.keywords,
                        'token_count': chunk.token_count,
                        'chunk_metadata': chunk.metadata
                    })
                
                total_tokens = sum(chunk.token_count for chunk in chunks)
                avg_tokens = total_tokens // len(chunks) if chunks else 0
                logger.info(f"✅ {filename}: {len(result)} chunks ({total_tokens} tokens, promedio {avg_tokens}/chunk)")
                return result
            except Exception as e:
                logger.error(f"Error en chunker inteligente para {filename}: {e}")
                logger.info("Usando método tradicional como fallback...")
        
        # FALLBACK: Método tradicional
        try:
            doc = docx.Document(file_path)
            content = []
            current_section = ""
            logger.info(f"Extrayendo (método tradicional): {filename}")

            # Párrafos
            for i, p in enumerate(doc.paragraphs):
                text = p.text.strip()
                if not text or len(text) < 5:
                    continue

                is_header = (
                    p.style.name.lower() in ['heading 1', 'heading 2', 'heading 3', 'título'] or
                    any(run.bold for run in p.runs) or
                    text.isupper() or '---' in text or 'Circle' in text or 'Target' in text
                )

                if is_header:
                    current_section = text
                else:
                    content.append({
                        'text': text,
                        'section': current_section,
                        'style': p.style.name,
                        'is_structured': self._is_structured_content(text),
                        'page_reference': f"doc_{i}"
                    })

            # Tablas
            for idx, table in enumerate(doc.tables):
                content.extend(self._extract_table_content(table, idx))

            structured = self._structure_for_rag(content, filename)
            logger.info(f"{filename}: {len(structured)} fragmentos útiles")
            return structured

        except Exception as e:
            logger.error(f"Error en {file_path}: {e}")
            return []

    def _extract_table_content(self, table, index: int) -> List[Dict]:
        rows = []
        try:
            for r_idx, row in enumerate(table.rows):
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells and len(' '.join(cells)) > 10:
                    rows.append({
                        'text': ' | '.join(cells),
                        'section': f'Tabla_{index + 1}',
                        'style': 'Table',
                        'is_structured': True,
                        'page_reference': f"table_{index}_{r_idx}"
                    })
        except Exception as e:
            logger.warning(f"Error en tabla: {e}")
        return rows

    def _is_structured_content(self, text: str) -> bool:
        patterns = [
            r'^\d+\.', r'^•', r'^- ', r'^\[', r'paso \d+', r'requisito', r'horario',
            r'lunes|martes|miércoles|jueves|viernes|sábado|domingo'
        ]
        return any(re.search(p, text.lower()) for p in patterns)
    
    def _is_relevant_content(self, text: str) -> bool:
        """Verifica si el contenido es relevante para indexación"""
        if not text or len(text.strip()) < 20:
            return False
        
        # Filtrar contenido irrelevante
        irrelevant_patterns = [
            r'^(página|page)\s*\d+',
            r'^tabla de contenido',
            r'^índice',
            r'^\s*\d+\s*$',
            r'^copyright|©',
            r'^todos los derechos reservados',
            r'^\s*\.{3,}',  # Puntos suspensivos
            r'^\s*_{3,}',   # Líneas de subrayado
            r'^\s*-{3,}',   # Líneas de guión
        ]
        
        text_lower = text.lower().strip()
        return not any(re.search(pattern, text_lower) for pattern in irrelevant_patterns)

    def _structure_for_rag(self, items: List[Dict], filename: str) -> List[Dict]:
        result = []
        base_cat = self._detect_category_from_filename(filename)

        for item in items:
            if not self._is_relevant_content(item['text']):
                continue

            cat = self._detect_category_from_content(item['text']) or base_cat
            text = self._format_for_rag(item['text'], item['section'], item['is_structured'])

            result.append({
                'content': text,
                'category': cat,
                'source': filename,
                'type': 'document_extract',
                'section': item['section'],
                'is_structured': item['is_structured']
            })
        return result

    def extract_from_txt(self, file_path: str) -> List[Dict[str, str]]:
        """Procesa archivos TXT planos"""
        try:
            filename = os.path.basename(file_path)
            logger.info(f"Extrayendo TXT: {filename}")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if len(content.strip()) < 50:
                logger.warning(f"Archivo TXT muy pequeño: {filename}")
                return []
            
            # Dividir en secciones por títulos o separadores
            sections = self._split_txt_into_sections(content)
            
            result = []
            for i, section in enumerate(sections):
                if len(section['text'].strip()) > 100:  # Mínimo 100 caracteres por sección
                    result.append({
                        'text': section['text'],
                        'section': section['title'] or f'Sección_{i+1}',
                        'style': 'Text',
                        'is_structured': section['is_structured'],
                        'page_reference': f'section_{i}'
                    })
            
            logger.info(f"TXT {filename}: {len(result)} secciones extraídas")
            return result
            
        except Exception as e:
            logger.error(f"Error procesando TXT {file_path}: {e}")
            return []
    
    def extract_from_pdf(self, file_path: str) -> List[Dict[str, str]]:
        """Procesa archivos PDF usando pdfplumber"""
        if not PDF_AVAILABLE:
            logger.error("pdfplumber no disponible para procesar PDF")
            return []
        
        try:
            filename = os.path.basename(file_path)
            logger.info(f"Extrayendo PDF: {filename}")
            
            content = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and len(text.strip()) > 50:
                        content.append({
                            'text': text,
                            'section': f'Página_{page_num + 1}',
                            'style': 'PDF_Page',
                            'is_structured': self._is_structured_content(text),
                            'page_reference': f'page_{page_num}'
                        })
            
            logger.info(f"PDF {filename}: {len(content)} páginas extraídas")
            return content
            
        except Exception as e:
            logger.error(f"Error procesando PDF {file_path}: {e}")
            return []
    
    def _split_txt_into_sections(self, content: str) -> List[Dict[str, Any]]:
        """Divide un archivo TXT en secciones lógicas"""
        sections = []
        lines = content.split('\n')
        current_section = []
        current_title = ""
        
        for line in lines:
            line = line.strip()
            
            # Detectar títulos/headers
            if self._is_txt_header(line):
                # Guardar sección anterior si existe
                if current_section:
                    text = '\n'.join(current_section).strip()
                    if len(text) > 50:
                        sections.append({
                            'title': current_title,
                            'text': text,
                            'is_structured': self._is_structured_content(text)
                        })
                
                # Iniciar nueva sección
                current_title = line
                current_section = []
            else:
                if line:  # Solo agregar líneas no vacías
                    current_section.append(line)
        
        # Agregar última sección
        if current_section:
            text = '\n'.join(current_section).strip()
            if len(text) > 50:
                sections.append({
                    'title': current_title,
                    'text': text,
                    'is_structured': self._is_structured_content(text)
                })
        
        # Si no se encontraron secciones, tratar todo como una sección
        if not sections and len(content.strip()) > 100:
            sections.append({
                'title': 'Contenido_Principal',
                'text': content.strip(),
                'is_structured': self._is_structured_content(content)
            })
        
        return sections
    
    def _is_txt_header(self, line: str) -> bool:
        """Detecta si una línea es un título o header"""
        if len(line) < 3:
            return False
        
        # Patrones de títulos
        patterns = [
            r'^#{1,6}\s',  # Markdown headers
            r'^\d+\.\s',   # Numerados
            r'^[A-Z][A-Z\s]{5,}$',  # TODO MAYÚSCULAS
            r'^\*\*.*\*\*$',  # **Título**
            r'^=+$|^-+$',   # Separadores
            r'^\s*\*\*?\s*[A-Z]',  # * TÍTULO o ** TÍTULO
        ]
        
        for pattern in patterns:
            if re.search(pattern, line):
                return True
        
        # Líneas cortas en mayúsculas
        if line.isupper() and len(line) < 60 and len(line) > 10:
            return True
        
        # Líneas que terminan con :
        if line.endswith(':') and len(line) < 80 and not line.startswith('http'):
            return True
        
        return False

    def _detect_category_from_filename(self, name: str) -> str:
        """Detecta categoría por nombre de archivo"""
        n = name.lower()
        mapping = {
            'deport': 'deportes', 'bienestar': 'bienestar_estudiantil', 'be': 'bienestar_estudiantil',
            'desarrollo': 'desarrollo_profesional', 'dl': 'desarrollo_profesional',
            'asuntos': 'asuntos_estudiantiles', 'tne': 'asuntos_estudiantiles', 'certificados': 'asuntos_estudiantiles',
            'calendario': 'academico', 'procedimientos': 'academico', 'manual': 'academico',
            'carreras': 'academico', 'directorio': 'academico', 'guia': 'general',
            'frecuentes': 'general', 'emergencia': 'seguridad', 'protocolo': 'seguridad'
        }
        for k, v in mapping.items():
            if k in n:
                return v
        return "general"

    def _detect_category_from_content(self, text: str) -> str:
        t = text.lower()
        if any(w in t for w in ['tne', 'tarjeta nacional estudiantil', 'pase escolar', 'certificado', 'seguro']):
            return 'asuntos_estudiantiles'
        if any(w in t for w in ['psicológico', 'salud mental', 'bienestar', 'crisis', 'línea ops', 'paedis']):
            return 'bienestar_estudiantil'
        if any(w in t for w in ['deporte', 'taller deportivo', 'gimnasio', 'caf', 'maiclub', 'entrenamiento']):
            return 'deportes'
        if any(w in t for w in ['trabajo', 'práctica', 'curriculum', 'bolsa trabajo', 'duoclaboral']):
            return 'punto_estudiantil'
        return "general"

    def _format_for_rag(self, text: str, section: str, structured: bool) -> str:
        text = re.sub(r'\s+', ' ', text).strip()
        if structured and section and len(section) > 5:
            return f"{section}: {text}"
        return text


class TrainingDataLoader:
    def __init__(self):
        self.data_loaded = False
        self.training_data_path = "./training_data"
        self.documents_path = "./app/documents"
        self.base_knowledge_loaded = False
        self.word_documents_loaded = False  # Ahora incluye DOCX, TXT y PDF
        self.document_processor = DocumentProcessor()

    def load_all_training_data(self):
        try:
            # OPTIMIZACIÓN: Solo cargar si no se ha cargado antes
            if self.data_loaded:
                logger.info("✅ Datos ya cargados (reutilizando)")
                return True
            
            logger.info("⚡ CARGA RÁPIDA INICIADA")

            # Solo cargar conocimiento base esencial (muy rápido)
            if not self.base_knowledge_loaded:
                self._load_corrected_base_knowledge()
                self.base_knowledge_loaded = True

            # OPTIMIZACIÓN: Cargar documentos TXT/DOCX solo si existen y es primera carga
            if not self.word_documents_loaded and os.path.exists(self.documents_path):
                # Contar archivos TXT y DOCX
                txt_count = len([f for f in os.listdir(self.documents_path) if f.endswith('.txt')])
                docx_count = len([f for f in os.listdir(self.documents_path) if f.endswith('.docx')])
                total_docs = txt_count + docx_count
                
                if total_docs > 0:
                    logger.info(f"📄 Cargando {txt_count} TXT + {docx_count} DOCX = {total_docs} documentos...")
                    self._load_documents()
                self.word_documents_loaded = True

            # OPTIMIZACIÓN: Cargar datos históricos de forma más eficiente
            self._load_historical_training_data()
            self._load_derivation_knowledge()
            
            # OPTIMIZACIÓN: Saltar cargas pesadas opcionales en startup
            # self._load_centro_ayuda_knowledge()  # Comentado: carga bajo demanda
            # self._load_specific_duoc_knowledge()  # Comentado: carga bajo demanda

            self.data_loaded = True
            logger.info("⚡ CARGA RÁPIDA FINALIZADA")
            return True
        except Exception as e:
            logger.error(f"Error en carga: {e}")
            return False

    def _load_corrected_base_knowledge(self):
        logger.info("Cargando conocimiento base corregido...")
        knowledge = [
            # TNE
            {"q": "Qué es TNE?", "a": "La TNE es la Tarjeta Nacional Estudiantil, beneficio para transporte público. Gestionada por JUNAEB. En Duoc UC se tramita en Punto Estudiantil.", "c": "asuntos_estudiantiles"},
            {"q": "tne duoc", "a": "Primera vez: $2.700. Revalidación: $1.100. Reposición: $3.600. Pago en caja o portal. Enviar comprobante a Puntoestudiantil_pnorte@duoc.cl", "c": "asuntos_estudiantiles"},
            {"q": "tarjeta nacional estudiantil", "a": "TNE = Tarjeta Nacional Estudiantil. Descuento en Metro, buses. Válida todo el año. Proceso vía JUNAEB, Duoc es intermediario.", "c": "asuntos_estudiantiles"},

            # DEPORTES
            {"q": "Gimnasio disponible?", "a": "Sí, gimnasio CAF en sede. Horario: L-V 13:00-20:20. Máximo 2 veces/semana. Inscripción en Punto Estudiantil.", "c": "deportes"},
            {"q": "talleres deportivos", "a": "Fútbol, voleibol, basquetbol, natación, boxeo, powerlifting, entrenamiento funcional. Gratuitos. Inscripciones semestrales.", "c": "deportes"},
            {"q": "gimnasio caf", "a": "CAF Duoc UC: Lunes a viernes 13:00-20:20. Sábado (por medio) 09:00-13:20. Uso con credencial estudiantil.", "c": "deportes"},

            # BIENESTAR
            {"q": "Apoyo psicológico", "a": "Línea OPS 24/7: +56 2 2820 3450. Sesiones virtuales: eventos.duoc.cl. Hasta 8 sesiones/año. Gratuito y confidencial.", "c": "bienestar_estudiantil"},
            {"q": "salud mental duoc", "a": "Apoyo psicológico virtual. Embajadores de Salud Mental. Talleres de bienestar. Contacto: avasquezm@duoc.cl", "c": "bienestar_estudiantil"},
            {"q": "línea ops", "a": "Urgencias emocionales 24/7: +56 2 2820 3450. Apoyo inmediato. Disponible fines de semana y festivos.", "c": "bienestar_estudiantil"},

            # DESARROLLO
            {"q": "bolsa de trabajo", "a": "duoclaboral.cl - Ofertas laborales y prácticas. Acceso con credenciales Duoc. Asesoría CV: ccortesn@duoc.cl", "c": "desarrollo_profesional"},
            {"q": "prácticas profesionales", "a": "Postulación desde 4to semestre. Plataforma: practicas.duoc.cl. Requisito: malla al día.", "c": "desarrollo_profesional"},
        ]

        for item in knowledge:
            self._add_to_rag(item["q"], item["a"], item["c"], "base", "original")
        logger.info(f"Base: {len(knowledge)} ítems")

    def _load_documents(self):
        """Carga documentos DOCX, TXT y PDF desde la carpeta documents/"""
        if not os.path.exists(self.documents_path):
            logger.warning("Carpeta documents/ no encontrada")
            return

        # Buscar todos los tipos de archivos soportados
        docx_files = glob.glob(os.path.join(self.documents_path, "*.docx"))
        txt_files = glob.glob(os.path.join(self.documents_path, "*.txt"))
        pdf_files = glob.glob(os.path.join(self.documents_path, "*.pdf"))
        
        total_files = len(docx_files) + len(txt_files) + len(pdf_files)
        print(f"\n📂 CARGANDO DOCUMENTOS:")
        print(f"   DOCX: {len(docx_files)} archivos")
        print(f"   TXT:  {len(txt_files)} archivos")
        print(f"   PDF:  {len(pdf_files)} archivos")
        print(f"   TOTAL: {total_files} archivos\n")
        logger.info(f"Documentos encontrados: {len(docx_files)} DOCX, {len(txt_files)} TXT, {len(pdf_files)} PDF")
        
        total_processed = 0
        total_chunks_added = 0

        # Procesar archivos DOCX
        if docx_files and DOCX_AVAILABLE:
            print(f"🔄 Procesando {len(docx_files)} documentos DOCX...")
            logger.info("Procesando documentos Word...")
            for path in docx_files:
                chunks_added = self._process_single_document(path, 'docx')
                total_processed += 1
                total_chunks_added += chunks_added
                print(f"   ✅ {os.path.basename(path)}: {chunks_added} chunks")
        elif docx_files and not DOCX_AVAILABLE:
            logger.error("Archivos .docx encontrados pero python-docx no está instalado")

        # Procesar archivos TXT
        if txt_files:
            print(f"\n🔄 Procesando {len(txt_files)} documentos TXT...")
            logger.info("Procesando documentos TXT...")
            for idx, path in enumerate(txt_files, 1):
                chunks_added = self._process_single_document(path, 'txt')
                total_processed += 1
                total_chunks_added += chunks_added
                print(f"   [{idx}/{len(txt_files)}] {os.path.basename(path)}: {chunks_added} chunks")

        # Procesar archivos PDF
        if pdf_files and PDF_AVAILABLE:
            print(f"\n🔄 Procesando {len(pdf_files)} documentos PDF...")
            logger.info("Procesando documentos PDF...")
            for path in pdf_files:
                chunks_added = self._process_single_document(path, 'pdf')
                total_processed += 1
                total_chunks_added += chunks_added
                print(f"   ✅ {os.path.basename(path)}: {chunks_added} chunks")
        elif pdf_files and not PDF_AVAILABLE:
            logger.error("Archivos .pdf encontrados pero pdfplumber no está instalado")

        print(f"\n✅ CARGA COMPLETADA:")
        print(f"   Archivos procesados: {total_processed}/{total_files}")
        print(f"   Chunks agregados: {total_chunks_added}\n")
        logger.info(f"TOTAL: {total_processed} archivos procesados, {total_chunks_added} chunks agregados")
        self.word_documents_loaded = True
    
    def _process_single_document(self, file_path: str, file_type: str) -> int:
        """Procesa un solo documento según su tipo"""
        name = os.path.basename(file_path)
        logger.info(f"Procesando {file_type.upper()}: {name}")
        
        try:
            # Extraer contenido según el tipo
            if file_type == 'docx':
                chunks = self.document_processor.extract_from_docx(file_path)
            elif file_type == 'txt':
                chunks = self.document_processor.extract_from_txt(file_path)
            elif file_type == 'pdf':
                chunks = self.document_processor.extract_from_pdf(file_path)
            else:
                logger.error(f"Tipo de archivo no soportado: {file_type}")
                return 0
            
            if not chunks:
                logger.warning(f"No se extrajo contenido de {name}")
                return 0

            # Procesar y agregar chunks al RAG
            added = 0
            for chunk in chunks:
                # Usar 'text' para TXT/PDF y 'content' para DOCX
                text_content = chunk.get('text') or chunk.get('content', '')
                if not text_content:
                    continue
                    
                enhanced = rag_engine.enhanced_normalize_text(text_content)
                
                # Detectar categoría del archivo
                category = self._detect_category_from_filename(name)
                if not category:
                    category = self._detect_category_from_content(text_content)
                
                # NUEVO: Pasar todos los metadatos enriquecidos del chunk
                chunk_metadata = chunk.get('chunk_metadata', {})
                if self._add_document_direct(enhanced, {
                    "type": f"document_{file_type}",
                    "category": category,
                    "source": name,
                    "section": chunk.get('section', ''),
                    "is_structured": chunk.get('is_structured', False),
                    "file_type": file_type,
                    "optimized": "true",
                    # NUEVOS metadatos del chunker inteligente
                    "keywords": chunk.get('keywords', []),
                    "token_count": chunk.get('token_count', 0),
                    "chunk_id": chunk.get('page_reference', ''),
                    "title": chunk_metadata.get('title', chunk.get('section', '')),
                    "has_overlap": chunk_metadata.get('has_overlap', False),
                    "fecha_procesamiento": chunk_metadata.get('fecha_procesamiento', '2025-11-26')
                }):
                    added += 1
            
            logger.info(f"{name}: {added}/{len(chunks)} fragmentos agregados al RAG")
            return added
            
        except Exception as e:
            logger.error(f"Error procesando {name}: {e}")
            return 0

    def _add_document_direct(self, doc: str, meta: Dict = None) -> bool:
        """USO SEGURO: rag_engine.add_document() → Evita acceso directo a collection"""
        try:
            # NUEVO: Metadatos enriquecidos con toda la información del chunker
            enhanced_metadata = {
                "source": meta.get('source', 'unknown'),
                "category": meta.get('category', 'general'),
                "type": meta.get('type', 'general'),
                "optimized": meta.get('optimized', 'false'),
                "section": meta.get('section', ''),
                "is_structured": meta.get('is_structured', False),
                # NUEVOS metadatos enriquecidos
                "keywords": ','.join(meta.get('keywords', [])) if meta.get('keywords') else '',
                "token_count": meta.get('token_count', 0),
                "chunk_id": meta.get('chunk_id', ''),
                "title": meta.get('title', ''),
                "fecha_procesamiento": meta.get('fecha_procesamiento', datetime.now().isoformat()[:10]),
                "has_overlap": meta.get('has_overlap', False),
                "file_type": meta.get('file_type', 'unknown')
            }
            
            return rag_engine.add_document(
                document=doc,
                metadata=enhanced_metadata
            )
        except Exception as e:
            logger.error(f"Error añadiendo documento: {e}")
            return False

    def _load_historical_training_data(self):
        pattern = os.path.join(self.training_data_path, "training_data_*.json")
        files = glob.glob(pattern)
        if not files:
            logger.warning("No hay training_data_*.json")
            return

        questions = []
        for f in files:
            with open(f, 'r', encoding='utf-8') as file:
                data = json.load(file)
                if isinstance(data, list):
                    questions.extend(data)
                elif isinstance(data, dict) and 'questions' in data:
                    questions.extend(data['questions'])

        for item in questions:
            q = item.get('input') or item.get('question', '')
            c = item.get('category', 'general')
            if q and len(q) > 5:
                self._add_document_direct(q, {"type": "historical", "category": c, "source": "training"})

    def _load_derivation_knowledge(self):
        items = [
            "DERIVACIÓN: Problemas técnicos → Centro de Ayuda: https://centroayuda.duoc.cl",
            "DERIVACIÓN: Consultas académicas → Jefatura de carrera"
        ]
        for doc in items:
            self._add_document_direct(doc, {"type": "derivacion", "category": "derivacion", "source": "system"})

    def _load_centro_ayuda_knowledge(self):
        docs = [
            "Centro de Ayuda Duoc UC: https://centroayuda.duoc.cl - Soporte técnico",
            "Portal del Estudiante: https://portal.duoc.cl - Acceso con RUT"
        ]
        for doc in docs:
            self._add_document_direct(doc, {"type": "info", "category": "general", "source": "centro_ayuda"})

    def _load_specific_duoc_knowledge(self):
        items = [
            {"doc": "UBICACIÓN: Complejo Maiclub (fútbol), Gimnasio Entretiempo (voleibol), Piscina Acquatiempo (natación)", "cat": "deportes"},
            {"doc": "CONTACTO: Claudia Cortés - ccortesn@duoc.cl - Desarrollo Laboral", "cat": "desarrollo_profesional"},
            {"doc": "CONTACTO: Adriana Vásquez - avasquezm@duoc.cl - Bienestar Estudiantil", "cat": "bienestar_estudiantil"}
        ]
        for i in items:
            self._add_document_direct(i["doc"], {"type": "contact", "category": i["cat"], "source": "duoc"})

    def _add_to_rag(self, q: str, a: str, cat: str, src: str, typ: str):
        doc = f"Pregunta: {q}\nRespuesta: {a}"
        enhanced = rag_engine.enhanced_normalize_text(doc)
        self._add_document_direct(enhanced, {
            "type": "faq",
            "category": cat,
            "source": src,
            "variation_type": typ,
            "optimized": "true"
        })

    def generate_knowledge_from_patterns(self):
        # OPTIMIZACIÓN: Solo generar si no se ha hecho antes
        if hasattr(self, '_patterns_generated') and self._patterns_generated:
            return
        
        patterns = [
            "Punto Estudiantil Plaza Norte: Santa Elena de Huechuraba 1660. L-V 8:30-19:00",
            "Certificado alumno regular: Digital gratis (portal), Impreso $1.000 (24h)",
            "Portal Estudiante: https://portal.duoc.cl",
            "Duoc Laboral: https://duoclaboral.cl"
        ]
        for doc in patterns:
            self._add_document_direct(doc, {"type": "pattern", "category": "general", "source": "generated"})
        
        self._patterns_generated = True

    def get_loading_status(self) -> Dict:
        return {
            "base_knowledge_loaded": self.base_knowledge_loaded,
            "documents_loaded": self.word_documents_loaded,
            "data_loaded": self.data_loaded,
            "docx_support": DOCX_AVAILABLE,
            "pdf_support": PDF_AVAILABLE
        }

    def _detect_category_from_filename(self, name: str) -> str:
        """Detecta categoría por nombre de archivo"""
        n = name.lower()
        mapping = {
            'deport': 'deportes', 'bienestar': 'bienestar_estudiantil', 'be': 'bienestar_estudiantil',
            'desarrollo': 'desarrollo_profesional', 'dl': 'desarrollo_profesional',
            'asuntos': 'asuntos_estudiantiles', 'tne': 'asuntos_estudiantiles', 'certificados': 'asuntos_estudiantiles',
            'calendario': 'academico', 'procedimientos': 'academico', 'manual': 'academico',
            'carreras': 'academico', 'directorio': 'academico', 'guia': 'general',
            'frecuentes': 'general', 'emergencia': 'seguridad', 'protocolo': 'seguridad'
        }
        for k, v in mapping.items():
            if k in n:
                return v
        return "general"
    
    def _detect_category_from_content(self, text: str) -> str:
        """Detecta categoría por contenido del texto"""
        t = text.lower()
        if any(w in t for w in ['tne', 'tarjeta nacional estudiantil', 'pase escolar', 'certificado', 'seguro']):
            return 'asuntos_estudiantiles'
        if any(w in t for w in ['psicológico', 'salud mental', 'bienestar', 'crisis', 'línea ops', 'paedis']):
            return 'bienestar_estudiantil'
        if any(w in t for w in ['deporte', 'taller deportivo', 'gimnasio', 'caf', 'maiclub', 'entrenamiento']):
            return 'deportes'
        if any(w in t for w in ['trabajo', 'práctica', 'curriculum', 'bolsa trabajo', 'duoclaboral']):
            return 'desarrollo_profesional'
        if any(w in t for w in ['calendario', 'semestre', 'evaluación', 'examen', 'matrícula']):
            return 'academico'
        if any(w in t for w in ['emergencia', 'evacuación', 'seguridad', 'protocolo', 'incendio']):
            return 'seguridad'
        return ""


# ========================================
# INSTANCIA GLOBAL OBLIGATORIA
# ========================================
training_loader = TrainingDataLoader()