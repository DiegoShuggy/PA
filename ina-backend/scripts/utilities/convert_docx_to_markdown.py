"""
Script de conversión DOCX → Markdown con frontmatter YAML
Duoc UC Plaza Norte - Sistema InA
Fecha: 01 Diciembre 2025

Convierte documentos DOCX a Markdown preservando:
- Estructura jerárquica (headers H1, H2, H3)
- Listas numeradas y con viñetas
- Formato de texto (negrita, cursiva)
- Tablas básicas
- Metadata enriquecida con frontmatter YAML
"""

import docx
from pathlib import Path
import re
from datetime import datetime
import yaml
import logging
from typing import Dict, List, Optional

logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)


class DOCXToMarkdownConverter:
    """Conversor inteligente de DOCX a Markdown con metadata enriquecida"""
    
    def __init__(self, metadata_mapping_path: str = None):
        """
        Args:
            metadata_mapping_path: Ruta al archivo categoria_mapping.yaml
        """
        self.metadata_mapping = {}
        
        if metadata_mapping_path and Path(metadata_mapping_path).exists():
            with open(metadata_mapping_path, 'r', encoding='utf-8') as f:
                self.metadata_mapping = yaml.safe_load(f)
            logger.info(f"✅ Cargado mapeo de {len(self.metadata_mapping)} categorías")
        else:
            logger.warning("⚠️ No se encontró categoria_mapping.yaml, usando metadata básica")
    
    def detect_category(self, text_content: str, filename: str) -> str:
        """
        Detecta la categoría del documento basado en contenido y nombre de archivo
        
        Args:
            text_content: Contenido completo del documento
            filename: Nombre del archivo DOCX
            
        Returns:
            Categoría detectada (tne, certificados, deportes, etc.)
        """
        text_lower = text_content.lower()
        filename_lower = filename.lower()
        
        # Mapeo de keywords a categorías
        category_keywords = {
            'tne': ['tne', 'tarjeta nacional', 'metro', 'transporte', 'bus'],
            'certificados': ['certificado', 'alumno regular', 'documento', 'concentración de notas'],
            'deportes': ['gimnasio', 'caf', 'deporte', 'maiclub', 'fitness', 'entrenamiento'],
            'bienestar': ['psicológico', 'psicólogo', 'salud mental', 'bienestar', 'ops'],
            'biblioteca': ['biblioteca', 'préstamo', 'libro', 'recurso bibliográfico'],
            'becas': ['beca', 'junaeb', 'gratuidad', 'financiamiento', 'económico'],
            'practicas': ['práctica', 'duoclaboral', 'empleo', 'cv', 'laboral'],
            'matricula': ['matrícula', 'pago', 'arancel', 'inscripción'],
            'emergencia': ['emergencia', 'evacuación', 'primeros auxilios', 'seguridad'],
            'contacto': ['contacto', 'teléfono', 'correo', 'ubicación', 'dirección']
        }
        
        # Primero intentar por nombre de archivo
        for category, keywords in category_keywords.items():
            if any(kw in filename_lower for kw in keywords):
                logger.info(f"📂 Categoría detectada por filename: {category}")
                return category
        
        # Luego por contenido (conteo de keywords)
        category_scores = {}
        for category, keywords in category_keywords.items():
            score = sum(text_lower.count(kw) for kw in keywords)
            if score > 0:
                category_scores[category] = score
        
        if category_scores:
            best_category = max(category_scores, key=category_scores.get)
            logger.info(f"📝 Categoría detectada por contenido: {best_category} (score: {category_scores[best_category]})")
            return best_category
        
        logger.warning("⚠️ No se pudo detectar categoría específica, usando 'general'")
        return 'general'
    
    def generate_frontmatter(self, doc_path: Path, category: str, 
                           text_content: str) -> Dict:
        """
        Genera metadata frontmatter YAML enriquecida
        
        Args:
            doc_path: Ruta al archivo DOCX original
            category: Categoría detectada
            text_content: Contenido del documento
            
        Returns:
            Diccionario con metadata para frontmatter
        """
        # Metadata base
        frontmatter = {
            'id': f"{category}_{doc_path.stem}",
            'source': doc_path.name,
            'source_type': 'docx_converted',
            'categoria': category,
            'fecha_conversion': datetime.now().strftime('%Y-%m-%d'),
            'fecha_modificacion_original': datetime.fromtimestamp(
                doc_path.stat().st_mtime
            ).strftime('%Y-%m-%d')
        }
        
        # Enriquecer con metadata del mapeo si existe
        if category in self.metadata_mapping:
            mapping = self.metadata_mapping[category]
            frontmatter.update({
                'departamento': mapping.get('departamento', 'general'),
                'keywords': mapping.get('keywords_base', []),
                'prioridad': mapping.get('prioridad', 'media'),
                'tema': mapping.get('tema_principal', category),
                'tipo_contenido': mapping.get('tipo_contenido', 'informativo')
            })
        
        # Extraer título principal del documento
        lines = text_content.split('\n')
        for line in lines[:10]:  # Revisar primeras 10 líneas
            clean_line = line.strip()
            if clean_line and len(clean_line) < 100:
                frontmatter['titulo'] = clean_line
                break
        
        return frontmatter
    
    def convert_paragraph_to_markdown(self, para) -> str:
        """
        Convierte un párrafo de DOCX a Markdown preservando formato
        
        Args:
            para: Objeto Paragraph de python-docx
            
        Returns:
            String en formato Markdown
        """
        text = para.text.strip()
        if not text:
            return ""
        
        # Detectar nivel de heading por estilo
        style_name = para.style.name.lower()
        
        if 'heading 1' in style_name or 'título 1' in style_name:
            return f"# {text}\n"
        elif 'heading 2' in style_name or 'título 2' in style_name:
            return f"## {text}\n"
        elif 'heading 3' in style_name or 'título 3' in style_name:
            return f"### {text}\n"
        elif 'heading 4' in style_name or 'título 4' in style_name:
            return f"#### {text}\n"
        
        # Detectar si es título por formato (negrita + corto)
        all_bold = all(run.bold for run in para.runs if run.text.strip())
        if all_bold and len(text) < 80 and not text.endswith('.'):
            # Es un título informal, convertir a H2
            return f"## {text}\n"
        
        # Detectar listas
        if para.style.name.lower().startswith('list'):
            # Detectar si es lista numerada o con viñetas
            if any(char.isdigit() for char in text[:3]):
                # Ya tiene número, mantener
                return f"{text}\n"
            else:
                return f"- {text}\n"
        
        # Aplicar formato inline (negrita, cursiva)
        formatted_text = self._apply_inline_formatting(para)
        
        return f"{formatted_text}\n"
    
    def _apply_inline_formatting(self, para) -> str:
        """Aplica formato inline (negrita, cursiva) a nivel de runs"""
        result = []
        
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            # Aplicar negrita
            if run.bold:
                text = f"**{text}**"
            
            # Aplicar cursiva
            if run.italic:
                text = f"*{text}*"
            
            result.append(text)
        
        return ''.join(result)
    
    def convert_table_to_markdown(self, table) -> str:
        """
        Convierte una tabla de DOCX a tabla Markdown
        
        Args:
            table: Objeto Table de python-docx
            
        Returns:
            String con tabla en formato Markdown
        """
        md_lines = []
        
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append('| ' + ' | '.join(cells) + ' |')
            
            # Agregar separador después del header
            if i == 0:
                md_lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        
        return '\n'.join(md_lines) + '\n'
    
    def convert_docx_to_markdown(self, docx_path: Path, 
                                 output_dir: Path = None) -> Optional[Path]:
        """
        Convierte un archivo DOCX a Markdown con frontmatter
        
        Args:
            docx_path: Ruta al archivo DOCX
            output_dir: Directorio de salida (opcional)
            
        Returns:
            Ruta al archivo Markdown generado
        """
        try:
            logger.info(f"📄 Procesando: {docx_path.name}")
            
            # Cargar documento
            doc = docx.Document(docx_path)
            
            # Convertir a texto plano primero (para detectar categoría)
            text_content = '\n'.join([para.text for para in doc.paragraphs])
            
            # Detectar categoría
            category = self.detect_category(text_content, docx_path.stem)
            
            # Generar frontmatter
            frontmatter = self.generate_frontmatter(docx_path, category, text_content)
            
            # Convertir párrafos a Markdown
            markdown_lines = []
            
            for element in doc.element.body:
                # Párrafos
                if element.tag.endswith('p'):
                    for para in doc.paragraphs:
                        if para._element == element:
                            md_line = self.convert_paragraph_to_markdown(para)
                            if md_line:
                                markdown_lines.append(md_line)
                            break
                
                # Tablas
                elif element.tag.endswith('tbl'):
                    for table in doc.tables:
                        if table._element == element:
                            md_table = self.convert_table_to_markdown(table)
                            markdown_lines.append(md_table)
                            break
            
            # Ensamblar documento final
            md_content = self._assemble_markdown_document(frontmatter, markdown_lines)
            
            # Determinar ruta de salida
            if output_dir is None:
                output_dir = Path("data/markdown") / category
            
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / f"{docx_path.stem}.md"
            
            # Guardar archivo
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            logger.info(f"✅ Convertido: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"❌ Error convirtiendo {docx_path.name}: {e}")
            return None
    
    def _assemble_markdown_document(self, frontmatter: Dict, 
                                    content_lines: List[str]) -> str:
        """
        Ensambla el documento Markdown final con frontmatter YAML
        
        Args:
            frontmatter: Diccionario con metadata
            content_lines: Líneas de contenido en Markdown
            
        Returns:
            Documento Markdown completo
        """
        # Generar YAML frontmatter
        yaml_front = yaml.dump(frontmatter, allow_unicode=True, sort_keys=False)
        
        # Ensamblar documento
        parts = [
            "---",
            yaml_front.strip(),
            "---",
            "",
            ''.join(content_lines)
        ]
        
        return '\n'.join(parts)
    
    def convert_directory(self, input_dir: Path, output_base_dir: Path = None):
        """
        Convierte todos los archivos DOCX en un directorio
        
        Args:
            input_dir: Directorio con archivos DOCX
            output_base_dir: Directorio base para salida
        """
        if output_base_dir is None:
            output_base_dir = Path("data/markdown")
        
        docx_files = list(input_dir.glob("*.docx"))
        
        if not docx_files:
            logger.warning(f"⚠️ No se encontraron archivos DOCX en {input_dir}")
            return
        
        logger.info(f"🔄 Encontrados {len(docx_files)} archivos DOCX")
        
        converted_count = 0
        failed_count = 0
        
        for docx_file in docx_files:
            # Ignorar archivos temporales de Word
            if docx_file.name.startswith('~$'):
                continue
            
            result = self.convert_docx_to_markdown(docx_file, output_base_dir)
            
            if result:
                converted_count += 1
            else:
                failed_count += 1
        
        logger.info(f"\n{'='*60}")
        logger.info(f"📊 RESUMEN DE CONVERSIÓN:")
        logger.info(f"   ✅ Convertidos: {converted_count}")
        logger.info(f"   ❌ Fallidos: {failed_count}")
        logger.info(f"   📂 Salida: {output_base_dir}")
        logger.info(f"{'='*60}\n")


def main():
    """Función principal de conversión"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Convertir DOCX a Markdown con frontmatter')
    parser.add_argument(
        '--input-dir',
        type=str,
        default='app/documents',
        help='Directorio con archivos DOCX (default: app/documents)'
    )
    parser.add_argument(
        '--output-dir',
        type=str,
        default='data/markdown',
        help='Directorio de salida (default: data/markdown)'
    )
    parser.add_argument(
        '--metadata-mapping',
        type=str,
        default='config/metadata/categoria_mapping.yaml',
        help='Ruta al archivo de mapeo de categorías'
    )
    
    args = parser.parse_args()
    
    # Crear conversor
    converter = DOCXToMarkdownConverter(metadata_mapping_path=args.metadata_mapping)
    
    # Convertir directorio
    input_path = Path(args.input_dir)
    output_path = Path(args.output_dir)
    
    if not input_path.exists():
        logger.error(f"❌ No existe el directorio: {input_path}")
        return
    
    converter.convert_directory(input_path, output_path)


if __name__ == "__main__":
    main()
